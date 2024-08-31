import pytesseract
from PIL import Image, ImageDraw
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
from transformers import pipeline

pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

# Load the NER model (bigcode/starpii gated - you can switch models if needed)
pipe = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english")

def detect_and_redact_text(text):
    # Use the NER model to detect entities
    entities = pipe(text)
    sensitive_data = []
    masked_text = text

    # Iterate over entities and redact the sensitive text
    for entity in entities:
        entity_text = text[entity['start']:entity['end']]
        sensitive_data.append((entity_text, entity['entity']))
        masked_text = masked_text.replace(entity_text, '****')  # Replace PII with "****"

    return masked_text, sensitive_data

def process_image(image_path):
    # Load image and perform OCR
    image = Image.open(image_path)
    ocr_text = pytesseract.image_to_string(image)

    # Detect and redact sensitive data
    masked_text, sensitive_data = detect_and_redact_text(ocr_text)

    print("Detected Sensitive Data:", sensitive_data)
    draw = ImageDraw.Draw(image)
    boxes = pytesseract.image_to_boxes(image)

    for b in boxes.splitlines():
        b = b.split(' ')
        word = b[0]  # The text detected from OCR
        if any(word in s[0] for s in sensitive_data):
            x1, y1, x2, y2 = int(b[1]), int(b[2]), int(b[3]), int(b[4])

            # Adjust y-coordinates to fit the PIL image format (origin top-left)
            y1 = image.height - y1
            y2 = image.height - y2

            # Ensure bounding boxes are correctly ordered
            x1, x2 = sorted([x1, x2])
            y1, y2 = sorted([y1, y2], reverse=True)

            # Ensure the coordinates are within the image boundaries
            x1 = max(0, min(image.width, x1))
            x2 = max(0, min(image.width, x2))
            y1 = max(0, min(image.height, y1))
            y2 = max(0, min(image.height, y2))

            print(f"Drawing rectangle: ({x1}, {y1}), ({x2}, {y2})")
            # Only draw valid rectangles
            if y1 >= y2 and x1 <= x2:
                draw.rectangle([x1, y1, x2, y2], fill="black")

    # Save redacted image
    masked_image_path = 'redacted_' + image_path
    image.convert("RGB").save(masked_image_path, "JPEG")
    print(f"Redacted image saved as '{masked_image_path}'")

def process_pdf(pdf_path):
    # Path to poppler binaries
    poppler_path = r'C:\Release-24.07.0-0\poppler-24.07.0\Library\bin'  # Update this to your path

    # Convert PDF to images
    images = convert_from_path(pdf_path, poppler_path=poppler_path)

    for page_number, image in enumerate(images):
        # Save each page of the PDF as an image
        image_path = f'page_{page_number}.jpg'
        image.save(image_path, 'JPEG')

        # Process each image (redact sensitive data)
        process_image(image_path)

def process_docx(docx_path):
    # Load DOCX document
    doc = Document(docx_path)
    full_text = []

    # Extract text from DOCX
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    full_text = ' '.join(full_text)

    # Detect and redact sensitive data
    masked_text, sensitive_data = detect_and_redact_text(full_text)
    print("Detected Sensitive Data:", sensitive_data)

    # Replace the paragraphs with masked text
    for para in doc.paragraphs:
        para.text = masked_text
    
    # Save the redacted DOCX file
    redacted_docx_path = 'redacted_' + docx_path
    doc.save(redacted_docx_path)
    print(f"Redacted DOCX saved as '{redacted_docx_path}'")

# Example usage
image_path = 'image.jpg'
pdf_path = r'C:\sih\sample.pdf'
docx_path = 'document.docx'

# Uncomment to process specific files
# process_image(image_path)  # Process image
process_pdf(pdf_path)  # Process PDF
# process_docx(docx_path)  # Process DOCX
