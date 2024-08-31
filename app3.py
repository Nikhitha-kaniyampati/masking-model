import pytesseract
from PIL import Image, ImageDraw
import spacy
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
import re

pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

# Load the spaCy model for NER
nlp = spacy.load("en_core_web_sm")

# Regex patterns for structured PII (e.g., Aadhaar, PAN)
aadhaar_pattern = r'\d{4}\s\d{4}\s\d{4}'
pan_pattern = r'[A-Z]{5}\d{4}[A-Z]{1}'

def detect_and_redact_text(text):
    # Detect structured PII using regex
    sensitive_data = []
    masked_text = re.sub(aadhaar_pattern, '**** **** ****', text)
    masked_text = re.sub(pan_pattern, '*****1234*', masked_text)
    
    # Detect entities using NER
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ in ['PERSON', 'ORG', 'GPE', 'CARDINAL']:
            sensitive_data.append((ent.text, ent.label_))
            masked_text = masked_text.replace(ent.text, '****')
    
    return masked_text, sensitive_data

def process_image(image_path):
    # Load image and perform OCR
    image = Image.open(image_path)
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

    # Extract text and bounding boxes
    texts = data['text']
    bbox = [(data['left'][i], data['top'][i], data['width'][i], data['height'][i]) for i in range(len(texts))]
    
    # Combine all text into one string
    full_text = ' '.join(texts)
    
    # Detect and mask sensitive data
    masked_text, sensitive_data = detect_and_redact_text(full_text)
    
    # Print detected sensitive data
    print("Detected Sensitive Data:", sensitive_data)
    
    # Draw rectangles to redact sensitive data
    draw = ImageDraw.Draw(image)
    for (x, y, w, h), text in zip(bbox, texts):
        if text in full_text and text not in masked_text:
            draw.rectangle([x, y, x + w, y + h], fill="black")

    # Save redacted image
    image.save('redacted_image.jpg')
    print("Redacted image saved as 'redacted_image.jpg'")
    
def process_pdf(pdf_path):
    # Convert PDF to images
    images = convert_from_path(pdf_path)
    
    for page_number, image in enumerate(images):
        # Save image from PDF
        image_path = f'page_{page_number}.png'
        image.save(image_path, 'PNG')
        
        # Process each image (page)
        process_image(image_path)

def process_docx(docx_path):
    # Load DOCX document
    doc = Document(docx_path)
    full_text = []

    # Extract text from DOCX
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    full_text = ' '.join(full_text)
    
    # Detect and mask sensitive data
    masked_text, sensitive_data = detect_and_redact_text(full_text)
    
    # Print detected sensitive data
    print("Detected Sensitive Data:", sensitive_data)
    
    # Replace text in the document with masked text
    for para in doc.paragraphs:
        if para.text in full_text:
            para.text = para.text.replace(para.text, masked_text)
    
    # Save redacted DOCX
    doc.save('redacted_document.docx')
    print("Redacted document saved as 'redacted_document.docx'")

# Example usage:
image_path = 'image.jpg'
pdf_path = 'sample.pdf'
# docx_path = '/path/to/your/document.docx'

# Process an image file
process_image(image_path)

# Process a PDF file
process_pdf(pdf_path)

# Process a DOCX file
# process_docx(docx_path)
